import os
from typing import List
from docx.document import Document
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Cm
from docx.table import _Cell
from docx.table import Table
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.text.paragraph import Paragraph
import json

 
def set_col_widths(table: Table, widths):
    # widths = (Cm(4), Cm(8))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
   
   
def add_formations(table: Table, formations: List[dict]):    
    for row_idx, formation in enumerate(formations):        
        # Ecole
        run = table.cell(row_idx, 0).paragraphs[0].add_run(formation['établissement'])
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(39, 55, 130)
        # Date
        run = table.cell(row_idx, 0).paragraphs[0].add_run(f'\n{formation["date"]}\n')
        run.font.bold = False
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(39, 55, 130)
        # Description
        run = table.cell(row_idx, 1).paragraphs[0].add_run(formation['titre'])
        run.font.bold = False
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(39, 55, 130)
        # Add a new row
        table.add_row()
 
    set_col_widths(table, widths=(Cm(6), Cm(12)))
   
def add_competences(table: Table, competences: dict):
    row_idx = 0
    for competence in competences.keys():
        if competences[competence]:        
            # Key
            run = table.cell(row_idx, 0).paragraphs[0].add_run(f'{competence.capitalize()} \n')
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(39, 55, 130)
           
            # Value
            run = table.cell(row_idx, 1).paragraphs[0].add_run(f'{competences[competence]} \n')
            run.font.bold = False
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(39, 55, 130)
           
            # Add a new row
            table.add_row()
            row_idx += 1
 
    set_col_widths(table, widths=(Cm(6), Cm(12)))
   
def add_langues(table: Table, langues: dict):
    for row_idx, langue in enumerate(langues):        
        # Key
        run = table.cell(row_idx, 0).paragraphs[0].add_run(f'{langue["nom"]} \n')
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(39, 55, 130)
        # Value
        run = table.cell(row_idx, 1).paragraphs[0].add_run(f'{langue["niveau"]}')
        run.font.bold = False
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(39, 55, 130)
        # Add a new row
        table.add_row()
   
    set_col_widths(table, widths=(Cm(6), Cm(12)))
 
   
def add_header_experience(idx_row, table, entreprise, date, color):
    # Entreprise
    cell: _Cell = table.cell(idx_row, 0)
    bg_color = parse_xml(fr'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(bg_color)
    run = cell.paragraphs[0].add_run(entreprise.upper())
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(255, 255, 255)
   
    # date
    cell: _Cell = table.cell(idx_row, 1)
    bg_color =  parse_xml(fr'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(bg_color)
    run = cell.paragraphs[0].add_run(date)
    run.font.bold = False
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(255, 255, 255)
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
 
def add_content_experience(idx_row: int, table: Table, titre: str, description: str, env_technique: str):
    merged_cell: _Cell = table.cell(idx_row, 0).merge(table.cell(idx_row, 1))
   
    # Titre
    run = merged_cell.paragraphs[0].add_run(f'\n{titre} \n\n')
    run.font.name = 'Arial'
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(39, 55, 130)
   
    # Description
    run = merged_cell.paragraphs[0].add_run(f'{description}\n\n')
    run.font.name = 'Arial'
    run.font.bold = False
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(39, 55, 130)
 
   
    # Environemment technique
    run = merged_cell.paragraphs[0].add_run('Environnement technique: ')
    run.font.name = 'Arial'
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(39, 55, 130)
   
    run = merged_cell.paragraphs[0].add_run(f'{env_technique} \n\n')
    run.font.name = 'Arial'
    run.font.bold = False
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(39, 55, 130)
   
 
def add_experiences(table: Table, experiences: List[dict], color: str):
    idx_row = 0
    for experience in experiences:
        # add header projet
        add_header_experience(idx_row,
                          table,
                          experience["nom de l'entreprise"],
                          experience['date'],
                          color)
        table.add_row()
        idx_row += 1
 
        # add content experience
        add_content_experience(idx_row,
                           table,
                           experience['nom du poste occupé'],
                           experience['description'],
                           experience['environnement technique'])
        table.add_row()
        idx_row += 1
    set_col_widths(table, widths=(Cm(11), Cm(5)))
   
 
def add_certificats(table: Table, certificats):
    for row_idx, certificat in enumerate(certificats):        
        # Ecole
        run = table.cell(row_idx, 0).paragraphs[0].add_run(certificat['nom'])
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(39, 55, 130)
        # Date
        run = table.cell(row_idx, 0).paragraphs[0].add_run(f'\n{certificat["date"]}\n')
        run.font.bold = False
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(39, 55, 130)
        table.add_row()
 
    set_col_widths(table, widths=(Cm(6), Cm(12)))
 
def add_header_projet(idx_row, table: Table, titre: str, date: str, color: str):
    # Entreprise
    cell: _Cell = table.cell(idx_row, 0)
    bg_color = parse_xml(fr'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(bg_color)
    run = cell.paragraphs[0].add_run(titre.upper())
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(255, 255, 255)
   
    # date
    cell: _Cell = table.cell(idx_row, 1)
    bg_color =  parse_xml(fr'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(bg_color)
    run = cell.paragraphs[0].add_run(date)
    run.font.bold = False
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(255, 255, 255)
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
 
def add_content_projet(idx_row, table: Table, description):
    merged_cell: _Cell = table.cell(idx_row, 0).merge(table.cell(idx_row, 1))
   
    # Description
    run = merged_cell.paragraphs[0].add_run(f'\n{description}\n\n')
    run.font.name = 'Arial'
    run.font.bold = False
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(39, 55, 130)
   
 
def add_projets(table: Table, projets, color: str):
    idx_row = 0
    for projet in projets:
        # add header projet
        add_header_projet(
            idx_row,
            table,
            projet["titre"],
            projet['date'],
            color
        )
        table.add_row()
        idx_row += 1
 
        # add content projet
        add_content_projet(
            idx_row,
            table,
            projet['description']
        )
                           
        table.add_row()
        idx_row += 1
    set_col_widths(table, widths=(Cm(11), Cm(5)))
   
def add_domaines_intervention(table: Table, domaines_intervention):
    row_idx = 0
    for domaine in domaines_intervention.keys():
        if domaines_intervention[domaine]:        
            # Key
            run = table.cell(row_idx, 0).paragraphs[0].add_run(f'{domaine.capitalize()} \n')
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(39, 55, 130)
           
            # Value
            run = table.cell(row_idx, 1).paragraphs[0].add_run(f'{domaines_intervention[domaine]} \n')
            run.font.bold = False
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(39, 55, 130)
           
            # Add a new row
            table.add_row()
            row_idx += 1
 
    set_col_widths(table, widths=(Cm(6), Cm(12)))
 
def generate_dev_senior(tables, data, mapping):
    # Add Formations
    formations = data['Formations']
    tab_formations = tables[mapping['Formations']]
    add_formations(tab_formations, formations)
 
    # Add Certificats
    certificats = data['Certifications']
    tab_certificats = tables[mapping['Certifications']]
    add_certificats(tab_certificats, certificats)
   
    # Add Compétences techniques
    competences_techniques = data['Compétences techniques']
    tab_competences_techniques = tables[mapping['Compétences techniques']]
    add_competences(tab_competences_techniques, competences_techniques)
 
    # Add Compétences fonctionnelles
    competences_foncitionnelles = data['Compétences fonctionnelles']
    tab_competences_fonctionnelles = tables[mapping['Compétences fonctionnelles']]
    add_competences(tab_competences_fonctionnelles, competences_foncitionnelles)
 
    # Add Langues
    langues = data['Langues']
    tab_langues = tables[mapping['Langues']]
    add_langues(tab_langues, langues)
 
    # Add Expériences
    experiences = data['Expériences professionnelles']
    tab_experiences = tables[mapping['Expériences professionnelles']]
    add_experiences(tab_experiences, experiences, color="#008BD2")
 
def generate_dev_junior(tables, data, mapping):
    # Add Formations
    formations = data['Formations']
    tab_formations = tables[mapping['Formations']]
    add_formations(tab_formations, formations)
 
    # Add Certificats
    certificats = data['Certifications']
    tab_certificats = tables[mapping['Certifications']]
    add_certificats(tab_certificats, certificats)
   
    # Add Compétences techniques
    competences = data['Compétences techniques']
    tab_competences = tables[mapping['Compétences techniques']]
    add_competences(tab_competences, competences)
 
    # Add Langues
    langues = data['Langues']
    tab_langues = tables[mapping['Langues']]
    add_langues(tab_langues, langues)
 
    # Add Expériences
    experiences = data['Expériences professionnelles']
    tab_experiences = tables[mapping['Expériences professionnelles']]
    add_experiences(tab_experiences, experiences, color="#008BD2")
   
    # Add Projets
    projets = data['Projets']
    tab_projets = tables[mapping['Projets']]
    add_projets(tab_projets, projets, color="#008BD2")
 
def generate_moa(tables, data, mapping):
    # Add Formations
    formations = data['Formations']
    tab_formations = tables[mapping['Formations']]
    add_formations(tab_formations, formations)
   
    # Add Domaines d'intervetion
    domaines_intervention = data["Domaines d'intervention"]
    tab_domaines = tables[mapping["Domaines d'intervention"]]
    add_domaines_intervention(tab_domaines, domaines_intervention)
    
    # Add compétences techniques
    competences = data['Compétences techniques']
    tab_competences = tables[mapping['Compétences techniques']]
    add_competences(tab_competences, competences)
    
    # Add Compétences fonctionnelles
    competences_foncitionnelles = data['Compétences fonctionnelles']
    tab_competences_fonctionnelles = tables[mapping['Compétences fonctionnelles']]
    add_competences(tab_competences_fonctionnelles, competences_foncitionnelles)
    
    # Add Langues
    langues = data['Langues']
    tab_langues = tables[mapping['Langues']]
    add_langues(tab_langues, langues)
 
    # Add Expériences
    experiences = data['Expériences professionnelles']
    tab_experiences = tables[mapping['Expériences professionnelles']]
    add_experiences(tab_experiences, experiences, color="#008BD2")


def generate_data_junior(tables, data, mapping):
    # Add Formations
    formations = data['Formations']
    tab_formations = tables[mapping['Formations']]
    add_formations(tab_formations, formations)
 
    # Add Certificats
    certificats = data['Certifications']
    tab_certificats = tables[mapping['Certifications']]
    add_certificats(tab_certificats, certificats)
   
    # Add Compétences techniques
    competences = data['Compétences techniques']
    tab_competences = tables[mapping['Compétences techniques']]
    add_competences(tab_competences, competences)
 
    # Add Langues
    langues = data['Langues']
    tab_langues = tables[mapping['Langues']]
    add_langues(tab_langues, langues)
 
    # Add Expériences
    experiences = data['Expériences professionnelles']
    tab_experiences = tables[mapping['Expériences professionnelles']]
    add_experiences(tab_experiences, experiences, color="#95C11F")
   
    # Add Projets
    projets = data['Projets']
    tab_projets = tables[mapping['Projets']]
    add_projets(tab_projets, projets, color="#95C11F")
    
def generate_data_senior(tables, data, mapping):
    # Add Formations
    formations = data['Formations']
    tab_formations = tables[mapping['Formations']]
    add_formations(tab_formations, formations)
 
    # Add Certificats
    certificats = data['Certifications']
    tab_certificats = tables[mapping['Certifications']]
    add_certificats(tab_certificats, certificats)
   
    # Add Compétences techniques
    competences_techniques = data['Compétences techniques']
    tab_competences_techniques = tables[mapping['Compétences techniques']]
    add_competences(tab_competences_techniques, competences_techniques)
 
    # Add Compétences fonctionnelles
    competences_foncitionnelles = data['Compétences fonctionnelles']
    tab_competences_fonctionnelles = tables[mapping['Compétences fonctionnelles']]
    add_competences(tab_competences_fonctionnelles, competences_foncitionnelles)
 
    # Add Langues
    langues = data['Langues']
    tab_langues = tables[mapping['Langues']]
    add_langues(tab_langues, langues)
 
    # Add Expériences
    experiences = data['Expériences professionnelles']
    tab_experiences = tables[mapping['Expériences professionnelles']]
    add_experiences(tab_experiences, experiences, color="#95C11F")
    
def choose_template(tables, data, type_cv):
    mapping_dev_senior = {
        'Formations': 1,
        'Certifications': 3,
        'Compétences techniques': 5,
        'Compétences fonctionnelles': 6,
        'Langues': 7,
        'Expériences professionnelles': 9
    }
    
   
    mapping_dev_junior = {
        'Formations': 1,
        'Certifications': 3,
        'Compétences techniques': 5,
        'Langues': 6,
        'Expériences professionnelles': 8,
        'Projets': 10
    }
   
    mapping_moa = {
        "Formations": 1,
        "Domaines d'intervention": 3,
        "Compétences techniques": 6,
        "Compétences fonctionnelles": 5,
        "Langues": 7,
        "Expériences professionnelles": 9
    }
    
    mapping_data_junior = mapping_dev_junior
    
    mapping_data_senior = mapping_dev_senior
   
   
 
    if type_cv == '-dev-senior':
        # Set mapping
        mapping = mapping_dev_senior
        generate_dev_senior(tables, data, mapping)
       
       
    if type_cv == '-dev-junior':
        # Set mapping
        mapping = mapping_dev_junior
        generate_dev_junior(tables, data, mapping)
       
       
    if type_cv == '-moa':
        # Set mapping
        mapping = mapping_moa
        generate_moa(tables, data, mapping)
        
    if type_cv == '-data-junior':
        # Set mapping
        mapping = mapping_data_junior
        generate_data_junior(tables, data, mapping)
        
    if type_cv == '-data-senior':
        # Set mapping
        mapping = mapping_data_senior
        generate_data_senior(tables, data, mapping)
       

def set_name(paragraphs: Paragraph, data: dict) -> str:
    name = (data["Profil"]["prenom"][0] + data["Profil"]["nom"][:2]).upper()
    run = paragraphs[2].add_run(name)
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial Black'
    run.font.color.rgb = RGBColor(0, 139, 210)

def set_caption(paragraphs, data: dict):
    caption = data['caption']
    run = paragraphs[4].add_run(caption)
    run.font.name = 'Arial'
    run.font.bold = False
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(39, 55, 130)   
       
 
 
def generate_docx(json_path, type_cv):
   
    doc: Document = Document("assests/template"+ type_cv +".docx")
       
    tables = doc.tables
    paragraphs = doc.paragraphs
   
    with open(json_path, 'r', encoding="utf-8") as f:
        data = json.load(f)
    
    # Set name and surname
    set_name(paragraphs, data)
    # Set caption
    set_caption(paragraphs, data)
    # Select template
    choose_template(tables, data, type_cv)
 
    
    docx_name = f'NewCo_Partners_{data["Profil"]["prenom"]}_{data["Profil"]["nom"]}.docx'
    curr = os.getcwd()
    nom_repertoire = os.path.join(curr, "output")
    os.makedirs(nom_repertoire, exist_ok=True)
   
    doc.save(os.path.join(nom_repertoire, docx_name))
 
    return docx_name
 
 
def main():
    json_path = 'uploaded/HASSAN_KIBOU_CV_JANVIER_24.json'
    file_name = 'cv-data-senior.txt'
    generate_docx(json_path, file_name, type_cv='-data-senior')
   
if __name__ == '__main__':
    main()
